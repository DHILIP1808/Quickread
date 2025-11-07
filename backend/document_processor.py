import io
import zipfile
import PyPDF2
from docx import Document as DocxDocument
from openpyxl import load_workbook

class DocumentProcessor:
    """Process various document formats and extract text from memory (no file storage)"""
    
    @staticmethod
    def process_pdf(file_bytes: bytes) -> str:
        """Extract text from PDF bytes"""
        text = ""
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
        return text

    @staticmethod
    def process_txt(file_bytes: bytes) -> str:
        """Extract text from TXT bytes"""
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text = file_bytes.decode('latin-1')
            except Exception as e:
                raise Exception(f"Error decoding text file: {str(e)}")
        return text

    @staticmethod
    def process_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX bytes"""
        text = ""
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = DocxDocument(docx_file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
        return text

    @staticmethod
    def process_xlsx(file_bytes: bytes) -> str:
        """Extract text from XLSX bytes"""
        text = ""
        try:
            xlsx_file = io.BytesIO(file_bytes)
            workbook = load_workbook(xlsx_file)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n=== Sheet: {sheet_name} ===\n"
                for row in sheet.iter_rows(values_only=True):
                    text += " | ".join(str(cell) if cell is not None else "" for cell in row) + "\n"
        except Exception as e:
            raise Exception(f"Error processing XLSX: {str(e)}")
        return text

    @staticmethod
    def process_zip(file_bytes: bytes) -> dict:
        """Extract and process files from ZIP bytes"""
        contents = {}
        try:
            zip_file = io.BytesIO(file_bytes)
            with zipfile.ZipFile(zip_file, 'r') as zip_ref:
                for file_info in zip_ref.filelist:
                    # Skip directories
                    if file_info.filename.endswith('/'):
                        continue
                    
                    # Get file extension
                    file_ext = file_info.filename.lower().split('.')[-1] if '.' in file_info.filename else ''
                    
                    # Read file content from ZIP
                    file_content = zip_ref.read(file_info.filename)
                    
                    # Process based on extension
                    try:
                        if file_ext == 'pdf':
                            contents[file_info.filename] = DocumentProcessor.process_pdf(file_content)
                        elif file_ext == 'txt':
                            contents[file_info.filename] = DocumentProcessor.process_txt(file_content)
                        elif file_ext == 'docx':
                            contents[file_info.filename] = DocumentProcessor.process_docx(file_content)
                        elif file_ext == 'xlsx':
                            contents[file_info.filename] = DocumentProcessor.process_xlsx(file_content)
                        else:
                            # For unsupported formats, store as text representation
                            contents[file_info.filename] = f"[Unsupported file type: .{file_ext}]"
                    except Exception as e:
                        contents[file_info.filename] = f"[Error processing file: {str(e)}]"
                        
        except Exception as e:
            raise Exception(f"Error processing ZIP: {str(e)}")
        return contents

    @staticmethod
    def process_document(file_bytes: bytes, file_extension: str) -> dict:
        """Process document bytes based on file type"""
        if file_extension == '.pdf':
            text = DocumentProcessor.process_pdf(file_bytes)
            return {"content": text, "format": "pdf"}
        elif file_extension == '.txt':
            text = DocumentProcessor.process_txt(file_bytes)
            return {"content": text, "format": "txt"}
        elif file_extension == '.docx':
            text = DocumentProcessor.process_docx(file_bytes)
            return {"content": text, "format": "docx"}
        elif file_extension == '.xlsx':
            text = DocumentProcessor.process_xlsx(file_bytes)
            return {"content": text, "format": "xlsx"}
        elif file_extension == '.zip':
            contents = DocumentProcessor.process_zip(file_bytes)
            return {"content": contents, "format": "zip"}
        else:
            raise Exception(f"Unsupported file format: {file_extension}")